from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 页边距 ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# ── 默认字体 ─────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


# ── 辅助函数 ─────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return p


def heading2(doc, text):
    """阶段标题：如 第一阶段"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return p


def heading3(doc, text):
    """人名小节：如 谢中凯"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return p


def task_label(doc, text):
    """加粗任务点，左缩进一级"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p


def method(doc, text, level=0):
    """方法说明，灰色小字，缩进更深"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.5)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def divider(doc):
    """细横线分隔"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ════════════════════════════════════════════════════════════
# 标题页
# ════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(40)
title_p.paragraph_format.space_after = Pt(6)
tr = title_p.add_run("基于 YOLOv11 的草莓采摘机械臂系统")
tr.bold = True
tr.font.size = Pt(18)
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
tr.font.name = "黑体"
tr._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

sub1 = doc.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub1.paragraph_format.space_after = Pt(4)
s1r = sub1.add_run("项目工作报告与任务分工")
s1r.bold = True
s1r.font.size = Pt(15)
s1r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
s1r.font.name = "黑体"
s1r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.paragraph_format.space_after = Pt(40)
s2r = sub2.add_run("同济大学 SITP 项目")
s2r.font.size = Pt(11)
s2r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

divider(doc)

# ════════════════════════════════════════════════════════════
# 一、项目概述
# ════════════════════════════════════════════════════════════
heading1(doc, "一、项目概述")
body(
    doc,
    "本项目以草莓采摘为应用场景，研究基于 YOLOv11-Pose 的采摘点检测模型与机械臂自动采摘系统的集成方法，"
    "覆盖视觉感知、运动规划、仿真验证与系统联调四个方向。"
    "最终目标是在真实或仿真环境中，实现机械臂对成熟草莓的准确定位与采摘，并形成一篇具有创新点的研究论文。",
)

# ─── 成员总览表 ───────────────────────────────────────────────
doc.add_paragraph()
tbl = doc.add_table(rows=6, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths = [Cm(2.6), Cm(5.0), Cm(7.6)]
for row in tbl.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_widths[i]

for i, h in enumerate(["姓名", "负责方向", "核心交付物"]):
    cell = tbl.cell(0, i)
    set_cell_bg(cell, "2E74B5")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

overview_data = [
    ("谢中凯", "项目统筹 · 文献调研 · 论文撰写", "文献综述、实验报告、论文初稿"),
    (
        "陆坚其",
        "MuJoCo 仿真 · IK 求解 · 采摘顺序算法",
        "可运行仿真 Demo、IK 脚本、采摘队列代码",
    ),
    ("王政鈞", "硬件集成 · 系统联调", "推理链路打通记录、联调测试数据"),
    ("刘炳尧", "视觉模型优化（遮挡场景）· 数据扩充", "v15 模型权重、遮挡对比指标报告"),
    ("喻士煊", "嵌入式部署 · 电控硬件", "TensorRT 导出验证、电控方案文档"),
]
row_colors = ["F2F7FC", "FFFFFF", "F2F7FC", "FFFFFF", "F2F7FC"]
for ri, (name, direction, deliverable) in enumerate(overview_data):
    row = tbl.rows[ri + 1]
    for ci, text in enumerate([name, direction, deliverable]):
        cell = row.cells[ci]
        set_cell_bg(cell, row_colors[ri])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        if ci == 0:
            run.bold = True

doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# 二、第一阶段任务（第 1 周）
# ════════════════════════════════════════════════════════════
heading1(doc, "二、第一阶段任务（第 1 周）")
body(
    doc,
    "目标：打通「视觉推理->3D 坐标->仿真抓取」完整链路，第一周结束时有可演示的系统。"
    "本阶段全员协同，所有人的工作围绕让链路跑通展开。",
)

divider(doc)

# 谢中凯
heading3(doc, "谢中凯 — 项目统筹 · 演示材料")

task_label(doc, "① 跟进各模块进度，维护联调问题记录表")
method(
    doc, "建立一个共享文档（腾讯文档/飞书），记录每天联调遇到的问题、责任人和解决状态"
)
method(doc, "每天下午与各模块负责人同步进展，确保瓶颈被及时识别")

task_label(doc, "② 整理第一周可展示的演示说明材料")
method(doc, "撰写演示流程说明（2 页以内）：系统能做什么、演示步骤、关键指标")
method(doc, "收集联调过程截图或录屏，整理成演示素材包")

divider(doc)

# 陆坚其
heading3(doc, "陆坚其 — MuJoCo 仿真环境搭建 · 基础 IK 验证")

task_label(doc, "① 安装 MuJoCo，加载 SO-101 仿真模型")
method(doc, "执行：pip install mujoco")
method(
    doc,
    "从 GitHub 仓库 TheRobotStudio/SO-ARM100 的 Simulation/ 目录下载 SO-101 的 MJCF 文件",
)
method(
    doc,
    '运行 MuJoCo 查看器，确认模型加载、关节可动：python -c "import mujoco; print(mujoco.__version__)"',
)

task_label(doc, "② 实现基础逆运动学（IK）求解")
method(doc, "安装 ikpy 库：pip install ikpy")
method(
    doc, '用 URDF 文件构建运动链：chain = ikpy.chain.Chain.from_urdf_file("so101.urdf")'
)
method(
    doc, "给定目标坐标 (X, Y, Z)，调用 chain.inverse_kinematics([x, y, z]) 得到各关节角"
)
method(doc, "验证：将关节角送入 MuJoCo，观察末端执行器是否到达目标点")

task_label(doc, "③ 搭建最简仿真场景，完成接口对接")
method(doc, "在 MJCF 场景中添加一个圆柱体代替草莓，放置在固定位置")
method(
    doc, '与王政鈞约定接口格式：接收 dict {"x": float, "y": float, "z": float}，单位 mm'
)
method(doc, "确认仿真中末端能运动到视觉给出的目标坐标附近，记录误差")

divider(doc)

# 王政鈞
heading3(doc, "王政鈞 — 硬件安装 · 推理链路打通 · 接口对接")

task_label(doc, "① 完成 Orbbec Gemini 335 与机械臂的物理安装，进行手眼标定")
method(
    doc,
    "固定相机位置（建议相机视角俯视草莓生长区域），用棋盘格标定板采集 15–20 张不同姿态图片",
)
method(
    doc,
    "使用 OpenCV 函数：cv2.calibrateCamera 标定相机内参，cv2.solvePnP 求外参（相机到机械臂基座的变换矩阵 T）",
)
method(doc, "将变换矩阵 T 存入文件，供 inference_windows.py 加载使用")

task_label(doc, "② 验证 inference_windows.py 完整运行")
method(
    doc,
    "确认运行环境：conda activate yolov11，权重路径指向 runs_pose/strawberry_pose_v14/weights/best.pt",
)
method(doc, "检查输出是否包含成熟草莓的 3D 采摘点坐标，单位 mm，格式合理")
method(doc, "深度读取时使用采摘点周围 3×3 像素的中位深度，而非单点值，减少噪声影响")

task_label(doc, "③ 将视觉输出传递给仿真模块，完成第一次联调")
method(doc, "与陆坚其约定接口：将 (X, Y, Z) 以 Python dict 或 JSON 格式传递")
method(
    doc,
    "录屏或截图记录第一次完整链路运行（相机识别草莓 → 给出坐标 → 仿真臂运动到目标点）",
)

divider(doc)

# 刘炳尧
heading3(doc, "刘炳尧 — 标注质量检查 · 为第二阶段优化准备")

task_label(doc, "① 检查现有标注文件中遮挡样本的可见性标志（v 值）")
method(doc, "打开 strawberry_pose_dataset/labels/train/ 目录下若干 .txt 文件")
method(doc, "每行格式为：class cx cy w h kx ky v，检查第 8 列（v）")
method(doc, "v=0 表示该关键点不参与训练；v=1 表示遮挡但有坐标；v=2 表示可见")
method(doc, "统计 v=0 的样本数量，记录比例，作为第二阶段修正的依据")

task_label(doc, "② 当前阶段不更换模型权重，配合王政鈞使用 v14 完成联调")
method(
    doc, "如联调中发现明显的预测偏差规律（如果梗始终偏左），记录下来作为后续改进依据"
)
method(doc, "收集联调中推理失败的样本（截图），用于标注数据集扩充")

divider(doc)

# 喻士煊
heading3(doc, "喻士煊 — 舵机通信验证 · 电控接口准备")

task_label(doc, "① 确认 SO-ARM100 舵机通信协议，完成单轴控制测试")
method(
    doc,
    "参考 LeRobot 仓库：lerobot/configs/robot/so100.yaml，确认通信方式（串口/USB）和波特率",
)
method(
    doc, "在 Windows 上用串口调试工具或 Python serial 库，发送单个关节的目标角度指令"
)
method(doc, "记录指令格式、返回值结构、控制频率限制")

task_label(doc, "② 封装基础电控接口，供联调模块调用")
method(doc, "编写 arm_controller.py，对外暴露一个函数：move_joints(joint_angles: list)")
method(doc, "内部将关节角度转换为对应的舵机脉宽或原始指令字节，通过串口发送")
method(
    doc, "在 autodl_scripts/ 目录下测试：python arm_controller.py（执行预设的单次运动）"
)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# 三、第二阶段任务（第 2–3 周）
# ════════════════════════════════════════════════════════════
heading1(doc, "三、第二阶段任务（第 2–3 周）")
body(
    doc,
    "目标：在联调链路跑通的基础上，分方向深入推进——视觉模型针对遮挡场景优化，"
    "仿真完善采摘顺序算法，硬件完成实机测试，文献调研同步展开。",
)

divider(doc)

# 谢中凯
heading3(doc, "谢中凯 — 文献调研 · 创新点梳理")

task_label(doc, "① 系统检索遮挡关键点检测相关文献")
method(doc, "检索平台：Google Scholar、IEEE Xplore、arXiv")
method(
    doc,
    '关键词："strawberry keypoint detection occluded"、"occluded pose estimation"、'
    '"fruit picking sequence planning"',
)
method(
    doc,
    "重点精读：Parsa et al. (J. Field Robotics 2023)、Liu et al. (ECCV 2022, arXiv:2208.00090)、"
    "Dai et al. 3MSP2 (Computers & Electronics in Agriculture, 2024)",
)
method(doc, "用 Zotero 或 Endnote 管理文献库，保持引用格式统一")

task_label(doc, "② 梳理论文创新点，与刘炳尧对齐实验设计")
method(
    doc,
    "核心创新点方向：YOLO-Pose 框架下，通过 v=1 遮挡标注 + 合成叶片增强，"
    "提升草莓果梗采摘点在遮挡场景下的预测精度——该方向目前无专项论文，存在文献空缺",
)
method(
    doc,
    "明确实验对比组：v14（无遮挡训练） vs v15（含遮挡标注+增强），分别报告可见/遮挡样本的 OKS/PCK",
)

task_label(doc, "③ 撰写论文框架与引言草稿")
method(doc, "确定论文结构：引言 → 相关工作 → 方法 → 实验 → 结论")
method(doc, "完成引言和相关工作章节初稿，参考文献不少于 15 篇")

divider(doc)

# 陆坚其
heading3(doc, "陆坚其 — 仿真场景完善 · 采摘顺序算法")

task_label(doc, "① 完善 MuJoCo 仿真场景")
method(
    doc, "在 MJCF 中添加草莓植株模型（可用多个圆柱体+球体组合近似）、桌面、多颗草莓目标"
)
method(doc, "设置合理的摩擦系数和碰撞参数，确保抓取时草莓不会异常弹飞")

task_label(doc, "② 实现并测试采摘顺序规划算法")
method(doc, "Step 1 — 深度排序：对所有检测结果按 Z 坐标升序排列，优先采摘最近的草莓")
method(
    doc,
    "Step 2 — 空间聚类：以 5 cm 欧氏距离为阈值，将 3D 检测点聚类；"
    "类间用贪心最近邻决定访问顺序，类内按深度排序",
)
method(
    doc, "Step 3 — 可达性验证：每次采摘前调用 IK 求解，若无解则跳过该目标，取队列下一项"
)
method(doc, "Step 4 — 重检测循环：每摘完一颗，重新触发视觉检测，更新采摘队列")

task_label(doc, "③ 对比实验")
method(doc, "在仿真中分别运行：随机顺序 / 纯深度排序 / 聚类感知排序，各运行 20 次")
method(doc, "记录指标：采摘成功率、碰撞次数、完成全部目标的平均用时")

divider(doc)

# 王政鈞
heading3(doc, "王政鈞 — 实机联调测试 · 误差分析")

task_label(doc, "① 开展多轮实机采摘测试，记录结果")
method(doc, "每轮测试设置 5 颗草莓（或等效目标），记录采摘成功/失败、失败原因分类")
method(doc, "失败原因分类：标定误差 / 深度噪声 / IK 无解 / 舵机执行误差 / 视觉漏检")

task_label(doc, "② 排查误差来源，针对性改进")
method(
    doc,
    "标定误差：若定位系统偏差 > 5mm，重新采集标定板图片，增加采集姿态数量至 25 张以上",
)
method(
    doc,
    "深度噪声：在 inference_windows.py 中将单点深度改为 3×3 中位滤波，"
    "调用方式：np.median(depth_array[y-1:y+2, x-1:x+2])",
)
method(
    doc,
    "时序问题：确认 pipeline.wait_for_frames(timeout_ms=100) 已启用，"
    "避免色深帧不对齐导致 3D 坐标错位",
)

task_label(doc, "③ 与喻士煊联调电控模块，验证完整闭环")
method(doc, "确认 arm_controller.py 接收关节角后，舵机实际运动角度与指令角度误差 < 1°")
method(
    doc,
    "完成一次完整闭环录制：相机看到草莓 → 推理给坐标 → IK 求解 → 舵机执行 → 末端到达目标点附近",
)

divider(doc)

# 刘炳尧
heading3(doc, "刘炳尧 — 遮挡标注修正 · 数据增强 · v15 训练")

task_label(doc, "① 修正遮挡样本的 v 值标注")
method(
    doc,
    "对所有 v=0 的遮挡 grasp_point，根据草莓形状估算果梗连接位置，填入坐标并改为 v=1",
)
method(doc, "估算方法：果梗底部通常位于 BBox 上边沿向内约 5–10% 处，沿草莓中轴线方向")
method(
    doc,
    "修正后在 label 文件中每行第 8 列由 0 改为 1，坐标列填入估算值（归一化到图片尺寸）",
)

task_label(doc, "② 实现合成遮挡数据增强（Copy-Paste Leaf Augmentation）")
method(doc, "Step 1：从现有图片中手动抠取 50+ 张叶片区域，保存为 PNG（带透明通道）")
method(
    doc,
    "Step 2：在训练 DataLoader 中编写自定义 Transform，"
    "随机选取 1–2 张叶片图贴到草莓 BBox 上方区域",
)
method(
    doc,
    "Step 3：检查粘贴后叶片是否覆盖 grasp_point 坐标；若覆盖，自动将对应 v 值改为 1",
)
method(
    doc,
    "基线版：也可先用 albumentations.CoarseDropout 做矩形遮挡增强验证思路可行性，"
    "再实现叶片 Paste 的精细版本",
)

task_label(doc, "③ 训练 v15 模型，完成对比评估")
method(doc, "运行训练脚本：python autodl_scripts/train_pose_local.py")
method(doc, "结果自动保存至 runs_pose/strawberry_pose_v15/weights/best.pt")
method(
    doc,
    "在测试集上分别计算可见关键点和遮挡关键点的 OKS（Object Keypoint Similarity）和 PCK@0.2，"
    "与 v14 对比，写成一张对比表格",
)

divider(doc)

# 喻士煊
heading3(doc, "喻士煊 — TensorRT 部署评估 · 气吹方案设计")

task_label(doc, "① 完成 TensorRT FP16 导出与速度对比测试")
method(doc, "在当前 Windows 机器（RTX 4070 Super）上执行导出：")
method(doc, "    from ultralytics import YOLO", level=1)
method(
    doc,
    '    YOLO("best.pt").export(format="engine", half=True, imgsz=640, batch=1)',
    level=1,
)
method(
    doc, "分别记录 PyTorch 推理延迟和 TensorRT FP16 推理延迟（单次推理，取 100 次均值）"
)
method(
    doc,
    "记录加速比；若后续部署 Jetson，需在 Jetson 上重新导出（.engine 文件与设备绑定）",
)
method(doc, "注意：首次推理有预热延迟，需在计时前先运行一次 dummy forward")

task_label(doc, "② 气吹结构电气方案调研与设计")
method(doc, "调研微型气泵规格：建议 6–12V 直流，出口风压 10–20 kPa，流量 3–8 L/min")
method(
    doc,
    "设计电磁阀控制回路：上位机通过 GPIO（Jetson）或 USB-Serial 模块发送高低电平，"
    "控制电磁阀开关，接通或断开气路",
)
method(doc, "输出器件选型报告：列出气泵型号、电磁阀型号、电源模块、接线图")

task_label(doc, "③ 完善 arm_controller.py，支持联调调用")
method(doc, "增加异常处理：串口断开、指令超时的回退逻辑")
method(
    doc,
    '增加状态回调：执行完成后向王政鈞的联调主循环返回 {"status": "done" / "failed"}',
)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# 四、第三阶段任务（第 4 周起）
# ════════════════════════════════════════════════════════════
heading1(doc, "四、第三阶段任务（第 4 周起）")
body(doc, "目标：汇总实验数据，撰写论文，完成项目结题报告。")

divider(doc)

heading3(doc, "谢中凯 — 论文主导撰写")
task_label(doc, "① 汇总各模块实验数据，完成结果分析章节")
method(doc, "收集：v14 vs v15 模型对比指标表、仿真采摘顺序对比数据、端到端延迟数据")
method(doc, "撰写实验结果与分析章节，数据以表格和折线图呈现")
task_label(doc, "② 完成论文全文初稿，组织内部审阅")
method(doc, "结构：引言 → 相关工作 → 系统架构 → 视觉模块 → 运动规划 → 实验 → 结论")
method(doc, "组织一次内部审阅，各成员对自己负责的技术章节进行核实和补充")

divider(doc)

heading3(doc, "陆坚其 — 仿真实验数据整理")
task_label(doc, "① 整理采摘顺序对比实验的完整数据，输出供论文引用的表格和图")
method(doc, "表格格式：策略名 | 成功率 | 碰撞次数均值 | 完成用时均值")
task_label(doc, "② 撰写论文中运动规划章节的技术描述")

divider(doc)

heading3(doc, "王政鈞 — 实机测试数据整理")
task_label(doc, "① 整理实机测试的成功率统计和误差分析，输出供论文引用")
task_label(doc, "② 协助谢中凯完善系统架构图和联调流程图")

divider(doc)

heading3(doc, "刘炳尧 — 视觉实验数据整理 · 论文视觉章节撰写")
task_label(doc, "① 整理 v14 vs v15 对比实验结果，含可见/遮挡样本分组指标")
task_label(doc, "② 撰写论文中数据增强与遮挡训练方法的技术章节")

divider(doc)

heading3(doc, "喻士煊 — 器件文档整理 · 论文补充")
task_label(doc, "① 整理气吹结构选型报告，纳入论文附录或系统实现章节")
task_label(doc, "② 整理 TensorRT 部署对比数据，输出供论文引用")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# 五、已完成工作总结（现阶段工作报告）
# ════════════════════════════════════════════════════════════
heading1(doc, "五、已完成工作总结")
body(
    doc,
    "以下为项目启动至当前时间节点已完成的工作，各成员接手新任务时可直接基于以下成果继续推进，"
    "无需从头搭建。",
)

doc.add_paragraph()

# 完成工作表格
done_tbl = doc.add_table(rows=10, cols=3)
done_tbl.style = "Table Grid"
done_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

done_widths = [Cm(3.5), Cm(6.5), Cm(5.2)]
for row in done_tbl.rows:
    for i, cell in enumerate(row.cells):
        cell.width = done_widths[i]

for i, h in enumerate(["完成项目", "具体内容", "位置 / 说明"]):
    cell = done_tbl.cell(0, i)
    set_cell_bg(cell, "2E74B5")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

done_data = [
    (
        "CVAT 图像标注",
        "手工标注 200 张草莓图片，标注类别 ripe（矩形框）和 grasp_point（关键点），有效样本 194 张",
        "pose_annotation/annotations.xml",
    ),
    (
        "数据集划分",
        "155 张训练集 / 39 张验证集，按 8:2 比例随机划分",
        "strawberry_pose_dataset/images/train（val）/",
    ),
    (
        "XML → YOLO Pose 转换",
        "cvat_xml2yolo_pose.py 将 CVAT XML 标注转为 YOLO Pose 格式，MARGIN=0.6",
        "autodl_scripts/cvat_xml2yolo_pose.py",
    ),
    (
        "视觉模型训练 v14",
        "YOLOv11n-Pose，训练 100 epoch，最优在 epoch 76\n"
        "Box mAP50 = 0.927，Pose mAP50 = 0.803，Pose mAP50-95 = 0.774",
        "runs_pose/strawberry_pose_v14/weights/best.pt",
    ),
    (
        "Orbbec 相机推理接入",
        "inference_windows.py 完整实现：相机读帧 → YOLO 推理 → 深度对齐 → 2D 转 3D 坐标输出，"
        "硬件 D2C 对齐已验证",
        "autodl_scripts/inference_windows.py",
    ),
    (
        "DatasetNinja 数据转换",
        "datasetninja2yolo_pose.py 实现 DatasetNinja 格式转 YOLO Pose（备用，"
        "因 peduncle 语义与 grasp_point 不一致，暂不混入训练）",
        "autodl_scripts/datasetninja2yolo_pose.py",
    ),
    (
        "训练脚本与配置",
        "train_pose_local.py + strawberry_pose_local.yaml，可直接复现 v14 训练",
        "autodl_scripts/train_pose_local.py",
    ),
    (
        "几何过滤推理",
        "inference_windows.py 中加入几何过滤逻辑：采摘点必须位于 BBox 上方 75% 区域内，"
        "过滤明显偏差的预测",
        "inference_windows.py 内置",
    ),
    (
        "SO-101 仿真模型",
        "SO-ARM100 官方 MJCF 模型可用，pip install mujoco 后直接加载",
        "TheRobotStudio/SO-ARM100 → Simulation/",
    ),
]

done_colors = ["F2F7FC", "FFFFFF"] * 5
for ri, (item, content, loc) in enumerate(done_data):
    row = done_tbl.rows[ri + 1]
    bg = done_colors[ri % 2]
    for ci, text in enumerate([item, content, loc]):
        cell = row.cells[ci]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        if ci == 0:
            run.bold = True

doc.add_paragraph()

# ── 当前局限性说明 ─────────────────────────────────────────
body(doc, "当前主要局限：")
p1 = doc.add_paragraph()
p1.paragraph_format.left_indent = Cm(0.6)
p1.paragraph_format.space_after = Pt(2)
p1.add_run("（1）数据量有限：").bold = True
p1.add_run(
    "仅 194 张标注图片，模型泛化性存在边界，对差异较大的拍摄角度和光照条件表现有所下降。"
    "后续扩充数据或完善数据增强后可继续训练 v15。"
).font.size = Pt(11)

p2 = doc.add_paragraph()
p2.paragraph_format.left_indent = Cm(0.6)
p2.paragraph_format.space_after = Pt(2)
p2.add_run("（2）遮挡场景未专项训练：").bold = True
p2.add_run(
    "当前标注中遮挡 grasp_point 的 v 值未统一处理，"
    "模型对叶片遮挡场景的预测能力有限，是第二阶段的主要优化方向。"
).font.size = Pt(11)

p3 = doc.add_paragraph()
p3.paragraph_format.left_indent = Cm(0.6)
p3.paragraph_format.space_after = Pt(2)
p3.add_run("（3）仿真与实机尚未联通：").bold = True
p3.add_run(
    "推理脚本已输出 3D 坐标，但坐标到机械臂关节角的转换、"
    "舵机控制的完整链路尚待第一阶段联调打通。"
).font.size = Pt(11)

# ════════════════════════════════════════════════════════════
# 六、接口约定
# ════════════════════════════════════════════════════════════
heading1(doc, "六、模块接口约定")
body(doc, "各模块间交接的数据格式在此统一约定，避免联调时出现格式不匹配：")
doc.add_paragraph()

iface_tbl = doc.add_table(rows=5, cols=4)
iface_tbl.style = "Table Grid"
iface_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

iface_widths = [Cm(3.0), Cm(3.0), Cm(5.0), Cm(4.0)]
for row in iface_tbl.rows:
    for i, cell in enumerate(row.cells):
        cell.width = iface_widths[i]

for i, h in enumerate(["上游", "下游", "数据格式", "备注"]):
    cell = iface_tbl.cell(0, i)
    set_cell_bg(cell, "2E74B5")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

iface_data = [
    (
        "刘炳尧\n（视觉模型）",
        "王政鈞\n（联调）",
        "更新 best.pt 路径，在 inference_windows.py 顶部 MODEL_PATH 变量修改",
        "不涉及数据格式",
    ),
    (
        "王政鈞\n（视觉推理）",
        "陆坚其\n（IK 模块）",
        '{"x": float, "y": float, "z": float, "conf": float}  单位：mm',
        "conf 为检测置信度，IK 可据此过滤低置信度目标",
    ),
    (
        "陆坚其\n（IK 求解）",
        "喻士煊\n（电控）",
        "[j1, j2, j3, j4, j5, j6]  单位：度（°）",
        "列表长度固定为 6，对应 SO-101 的 6 个关节",
    ),
    (
        "喻士煊\n（电控）",
        "王政鈞\n（联调主循环）",
        '{"status": "done" | "failed", "error": str | None}',
        "failed 时 error 字段说明原因",
    ),
]

iface_colors = ["F2F7FC", "FFFFFF", "F2F7FC", "FFFFFF"]
for ri, row_data in enumerate(iface_data):
    for ci, text in enumerate(row_data):
        cell = iface_tbl.cell(ri + 1, ci)
        set_cell_bg(cell, iface_colors[ri])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10.5)

doc.add_paragraph()
body(doc, "协作约定：")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.6)
p.paragraph_format.space_after = Pt(2)
p.add_run("·  ").bold = True
p.add_run(
    "代码统一放在 autodl_scripts/ 目录下，文件名含义明确，不使用 test1.py 等临时命名"
)
p2 = doc.add_paragraph()
p2.paragraph_format.left_indent = Cm(0.6)
p2.paragraph_format.space_after = Pt(2)
p2.add_run("·  ").bold = True
p2.add_run("接口格式变更需提前告知下游负责人，并更新本文档接口约定表")
p3 = doc.add_paragraph()
p3.paragraph_format.left_indent = Cm(0.6)
p3.paragraph_format.space_after = Pt(2)
p3.add_run("·  ").bold = True
p3.add_run(
    "模型权重命名规则：strawberry_pose_v{版本号}/weights/best.pt，当前最优为 v14，下次训练存 v15"
)
p4 = doc.add_paragraph()
p4.paragraph_format.left_indent = Cm(0.6)
p4.paragraph_format.space_after = Pt(2)
p4.add_run("·  ").bold = True
p4.add_run("每阶段结束时各成员提交一份进展小结（半页以内），由谢中凯汇总后统一存档")

# ════════════════════════════════════════════════════════════
# 保存
# ════════════════════════════════════════════════════════════
out_path = r"C:\Users\xzKan\Desktop\机械臂草莓\项目任务分工_v2.docx"
doc.save(out_path)
print(f"文档已生成：{out_path}")
